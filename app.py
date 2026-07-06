import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from difflib import SequenceMatcher
import io

# -- UI KONFIGURATION --
st.set_page_config(page_title="Multi-Doc Fusion Pro", layout="wide")

st.markdown("""
    <style>
    .diff-box { padding: 10px; border-radius: 5px; border: 1px solid #ddd; height: 350px; overflow-y: auto; font-family: monospace; white-space: pre-wrap; font-size: 14px;}
    .add { background-color: #d4edda; text-decoration: underline; font-weight: bold; }
    .rem { background-color: #f8d7da; text-decoration: line-through; opacity: 0.7; }
    </style>
""", unsafe_allow_html=True)

# Farb-Palette für Word-Highlights (für bis zu 5 zusätzliche Versionen)
COLOR_PALETTE = [
    ("Gelb", WD_COLOR_INDEX.YELLOW, "#FFF3CD"),
    ("Türkis", WD_COLOR_INDEX.TURQUOISE, "#D1ECF1"),
    ("Pink", WD_COLOR_INDEX.PINK, "#F8D7DA"),
    ("Hellgrün", WD_COLOR_INDEX.BRIGHT_GREEN, "#D4EDDA"),
    ("Violett", WD_COLOR_INDEX.VIOLET, "#E2D9F3")
]

def reset_app():
    for key in list(st.session_state.keys()):
        del st.session_state[key]

def get_text(file):
    doc = docx.Document(file)
    return [p.text for p in doc.paragraphs if p.text.strip() != ""]

def create_multi_fused_docx(master_name, master_text, versions_dict):
    doc = docx.Document()
    
    # --- 1. LEGENDE ERSTELLEN ---
    title = doc.add_heading('Legende der Dokumentenfusion', level=1)
    p_intro = doc.add_paragraph("Dieses Dokument beinhaltet die zusammengeführten Inhalte aus mehreren Versionen.")
    
    # Master in der Legende
    p_master = doc.add_paragraph()
    r_master = p_master.add_run(f"• MASTER-DOKUMENT (Basis): ")
    r_master.bold = True
    p_master.add_run(master_name)
    
    # Versionen in der Legende
    for idx, (v_name, _) in enumerate(versions_dict.items()):
        color_name, color_idx, _ = COLOR_PALETTE[idx % len(COLOR_PALETTE)]
        p_ver = doc.add_paragraph()
        r_ver = p_ver.add_run(f"• ERGÄNZUNGEN AUS: {v_name} ")
        r_ver.bold = True
        r_highlight = p_ver.add_run(f" [ Markierung: {color_name} ] ")
        r_highlight.font.highlight_color = color_idx

    doc.add_page_break()
    
    # --- 2. MULTI-FUSION LOGIK ---
    doc.add_heading('Fusioniertes Dokumenten-Ergebnis', level=1)
    
    # Wir starten mit dem Master-Text als Basis
    # Und injizieren sequenziell die Ergänzungen der anderen Versionen
    current_text = list(master_text)
    
    for idx, (v_name, v_text) in enumerate(versions_dict.items()):
        color_name, color_idx, _ = COLOR_PALETTE[idx % len(COLOR_PALETTE)]
        matcher = SequenceMatcher(None, master_text, v_text)
        
        # Um das Word-Dokument sauber aufzubauen, schreiben wir den Basis-Text
        # und fügen an den erkannten Stellen die markierten Absätze ein.
        if idx == 0:
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    for line in master_text[i1:i2]:
                        doc.add_paragraph(line)
                elif tag in ('insert', 'replace'):
                    for line in v_text[j1:j2]:
                        p = doc.add_paragraph()
                        run = p.add_run(f"[{v_name}]: {line}")
                        run.font.highlight_color = color_idx
        else:
            # Für weitere Versionen hängen wir geänderte Absätze als separate Blöcke an die passenden Stellen
            # (In einer simplen, robusten Logik werden diese als Kommentare/Ergänzungen markiert)
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag in ('insert', 'replace'):
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{v_name}]: {line}")
                    run.font.highlight_color = color_idx

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# -- APP LAYOUT --
st.title("📑 Multi-Dokumenten Vergleich & Fusion")
st.write("Lade mehrere Versionen hoch, bestimme ein Master-Dokument und erstelle ein fusioniertes Word-File inklusive Farb-Legende.")

# Mehrfach-Upload erlauben
uploaded_files = st.file_uploader("Wähle 2 oder mehr Word-Dateien aus (.docx)", type="docx", accept_multiple_files=True)

if uploaded_files and len(uploaded_files) >= 2:
    st.divider()
    
    # Dateinamen für Auswahl extrahieren
    file_map = {file.name: file for file in uploaded_files}
    file_names = list(file_map.keys())
    
    # Master-Dokument auswählen
    col_sel1, col_sel2 = st.columns([1, 2])
    with col_sel1:
        master_name = st.selectbox("👑 Wähle das Master-Dokument (Basis):", file_names)
    
    master_file = file_map[master_name]
    master_text = get_text(master_file)
    
    # Alle anderen Versionen in ein Dict packen
    other_versions = {name: get_text(file) for name, file in file_map.items() if name != master_name}
    
    # --- VISUELLER VERGLEICH IN TABS ---
    st.subheader("🔍 Detail-Vergleich (Master vs. jeweilige Version)")
    
    # Erstelle für jede Vergleichsdatei einen eigenen Tab
    tab_names = [f"vs. {name}" for name in other_versions.keys()]
    tabs = st.tabs(tab_names)
    
    for idx, (v_name, v_text) in enumerate(other_versions.items()):
        with tabs[idx]:
            c1, c2 = st.columns(2)
            c1.markdown(f"**👑 Master: {master_name}**")
            c2.markdown(f"**📝 Version: {v_name}**")
            
            matcher = SequenceMatcher(None, master_text, v_text)
            diff_html_1, diff_html_2 = "", ""
            
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    chunk = "\n".join(master_text[i1:i2]) + "\n"
                    diff_html_1 += chunk
                    diff_html_2 += chunk
                elif tag == 'insert':
                    diff_html_2 += f'<span class="add">{" ".join(v_text[j1:j2])}</span>\n'
                elif tag == 'delete':
                    diff_html_1 += f'<span class="rem">{" ".join(master_text[i1:i2])}</span>\n'
                elif tag == 'replace':
                    diff_html_1 += f'<span class="rem">{" ".join(master_text[i1:i2])}</span>\n'
                    diff_html_2 += f'<span class="add">{" ".join(v_text[j1:j2])}</span>\n'

            c1.markdown(f'<div class="diff-box">{diff_html_1}</div>', unsafe_allow_html=True)
            c2.markdown(f'<div class="diff-box">{diff_html_2}</div>', unsafe_allow_html=True)

    # --- FUSION & DOWNLOAD ---
    st.divider()
    st.subheader("✨ Dokumenten-Fusion")
    st.write("Generiert ein neues Word-Dokument mit einer Legende auf der ersten Seite und fügt alle Ergänzungen farblich markiert zusammen.")
    
    fused_file = create_multi_fused_docx(master_name, master_text, other_versions)
    
    st.download_button(
        label="📥 Fusioniertes Multi-Dokument herunterladen (.docx)", 
        data=fused_file, 
        file_name="multi_fusion_result.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )

elif uploaded_files and len(uploaded_files) == 1:
    st.info("Bitte lade noch mindestens eine weitere Datei hoch, um den Vergleich zu starten.")

st.divider()
st.button("🔄 Alle Dateien löschen & neu starten", on_click=reset_app)
